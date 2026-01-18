#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""本地刷题软件（MVP，离线）

满足需求：
- 输入 Word/文档（优先 .docx；也支持 .doc/.html/.txt/.md）
- 自动识别题目并按“每题一页”刷题
- 支持翻页/跳题/随机
- 错题本：做错自动加入；错题练习中连续答对3次自动移除；答错则连续正确清零
- 练习记录：attempt 全本地保存

运行：
  python app.py

可选依赖：
  pip install python-docx

打包（可选）：
  pip install pyinstaller
  pyinstaller -F -w app.py
"""

from __future__ import annotations

import hashlib
import html as ihtml
import json
import os
import re
import sqlite3
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import tkinter as tk
from tkinter import ttk, filedialog, messagebox


APP_DIR = Path(__file__).resolve().parent
DB_PATH = APP_DIR / "quiz.db"


# -----------------------------
# Utils
# -----------------------------

def now_ts() -> str:
    return datetime.now().isoformat(timespec="seconds")


def sha1_text(s: str) -> str:
    return hashlib.sha1(s.encode("utf-8", errors="ignore")).hexdigest()


def normalize_spaces(s: str) -> str:
    s = re.sub(r"\r\n?", "\n", s)
    s = re.sub(r"[\t\f\v]", " ", s)
    s = re.sub(r"[ \u00A0]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return "\n".join(line.rstrip() for line in s.split("\n")).strip()


def norm_letters(s: str) -> str:
    letters = re.sub(r"[^A-H]", "", (s or "").upper())
    # 多选按集合比较
    return "".join(sorted(set(letters)))


def judge_to_bool(ans: str) -> Optional[bool]:
    a = (ans or "").strip()
    if not a:
        return None
    if a in ("对", "正确", "√", "是", "T", "TRUE", "True"):
        return True
    if a in ("错", "错误", "×", "否", "F", "FALSE", "False"):
        return False
    # 有些题库用 A/B 表示对/错
    letters = norm_letters(a)
    if letters == "A":
        return True
    if letters == "B":
        return False
    return None


# -----------------------------
# Document to plain text
# -----------------------------

def html_to_text(html: str) -> str:
    """Very small HTML->text converter (no external deps)."""
    html = re.sub(r"(?is)<\s*br\s*/?>", "\n", html)
    html = re.sub(r"(?is)<\s*/\s*p\s*>", "\n", html)
    html = re.sub(r"(?is)<\s*p\b[^>]*>", "", html)
    html = re.sub(r"(?is)<\s*/\s*tr\s*>", "\n", html)
    html = re.sub(r"(?is)<\s*/\s*td\s*>", "\t", html)
    html = re.sub(r"(?is)<\s*li\b[^>]*>", "\n- ", html)
    html = re.sub(r"(?is)<\s*/\s*li\s*>", "", html)
    html = re.sub(r"(?is)<[^>]+>", "", html)
    html = ihtml.unescape(html)
    return normalize_spaces(html)


def extract_plain_text(file_path: Path) -> str:
    suf = file_path.suffix.lower()

    if suf in (".txt", ".md"):
        return normalize_spaces(file_path.read_text("utf-8", errors="ignore"))

    if suf in (".html", ".htm"):
        html = file_path.read_text("utf-8", errors="ignore")
        # 有 bs4 就用（更能保留换行结构），否则用轻量替代
        try:
            from bs4 import BeautifulSoup  # type: ignore

            soup = BeautifulSoup(html, "html.parser")
            return normalize_spaces(soup.get_text("\n"))
        except Exception:
            return html_to_text(html)

    if suf == ".docx":
        try:
            import docx  # type: ignore
        except Exception as e:
            raise RuntimeError("解析 .docx 需要安装 python-docx：pip install python-docx") from e
        doc = docx.Document(str(file_path))
        parts: List[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        # tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        row_text.append(t)
                if row_text:
                    parts.append("\t".join(row_text))
        return normalize_spaces("\n".join(parts))

    if suf == ".doc":
        data = file_path.read_bytes()
        # 这类题库 doc 往往是“Word容器 + HTML片段”
        text = data.decode("utf-8", errors="ignore")
        m = re.search(r"<!DOCTYPE\s+html", text, flags=re.IGNORECASE)
        if m:
            html = text[m.start():]
        else:
            idx = text.lower().find("<html")
            html = text[idx:] if idx != -1 else text
        try:
            from bs4 import BeautifulSoup  # type: ignore

            soup = BeautifulSoup(html, "html.parser")
            return normalize_spaces(soup.get_text("\n"))
        except Exception:
            return html_to_text(html)

    # fallback: best-effort decode
    return normalize_spaces(file_path.read_text("utf-8", errors="ignore"))


# -----------------------------
# Parsing rules
# -----------------------------

"""章节识别一定要“像标题”，避免把题干里的“[简答题]”之类误判为章节切换。

策略：
1) 优先匹配“整行标题”（行首 ^，且该行基本只有标题文本）
2) 兼容“二、判断题”这种同一行写法
"""

SECTION_PATTERNS: List[Tuple[re.Pattern, str]] = [
    (re.compile(r"(?m)^\s*(?:[一二三四五六七八九十\d]+[、\.]\s*)?(?:单项选择题|单选题)\s*$"), "单选题"),
    (re.compile(r"(?m)^\s*(?:[一二三四五六七八九十\d]+[、\.]\s*)?(?:多项选择题|多选题)\s*$"), "多选题"),
    (re.compile(r"(?m)^\s*(?:[一二三四五六七八九十\d]+[、\.]\s*)?判断题\s*$"), "判断题"),
    (re.compile(r"(?m)^\s*(?:[一二三四五六七八九十\d]+[、\.]\s*)?填空题\s*$"), "填空题"),
    (re.compile(r"(?m)^\s*(?:[一二三四五六七八九十\d]+[、\.]\s*)?阅读理解\s*$"), "阅读理解"),
    (re.compile(r"(?m)^\s*(?:[一二三四五六七八九十\d]+[、\.]\s*)?简答题\s*$"), "简答题"),
]


def build_section_markers(text: str) -> List[Tuple[int, str]]:
    markers: List[Tuple[int, str]] = []
    for pat, name in SECTION_PATTERNS:
        for m in pat.finditer(text):
            markers.append((m.start(), name))
    markers.sort(key=lambda x: x[0])
    return markers


def section_at(markers: List[Tuple[int, str]], pos: int) -> str:
    cur = "未知"
    for p, name in markers:
        if p <= pos:
            cur = name
        else:
            break
    return cur


# 题号起始：优先“1、”这种中文分隔。若是“1.”则要求点号后面不是数字，
# 以避免把“1.2.1.*”这类网络前缀/版本号误拆成题目。
QSTART = re.compile(r"(?m)^\s*(\d+)(?:、|\.(?!\d))\s*")
OPT_SINGLE = re.compile(r"^([A-H])[、\.)]$")
OPT_INLINE = re.compile(r"^([A-H])[、\.)]\s*(.+)$")
ANS_MARK = re.compile(r"^(正确答案|参考答案|答案)\s*[:：]?")
ANA_MARK = re.compile(r"^(解析|答案解析)\s*[:：]?")
# 阅读理解子问通常是“(1)、”“（2）...”之类。
# 注意：题干里经常会出现“(20.0)”这种分值行，因此这里要求后面紧跟分隔符（、/./）来避免误判。
SUBQ_START = re.compile(r"(?m)^\s*[（\(](\d{1,2})[）\)]\s*[、\.)]")


def parse_answer_from_lines(lines: List[str], start_idx: int) -> Tuple[str, int]:
    """Find answer marker from lines[start_idx:], return (answer, next_index)."""
    i = start_idx
    while i < len(lines):
        l = lines[i]
        if ANS_MARK.match(l):
            # same-line answer
            parts = re.split(r"[:：]", l, maxsplit=1)
            if len(parts) == 2 and parts[1].strip():
                return parts[1].strip(), i + 1
            # next-line answer
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            return (lines[j].strip() if j < len(lines) else ""), j + 1
        i += 1
    return "", len(lines)


def parse_analysis_from_lines(lines: List[str], start_idx: int) -> str:
    i = start_idx
    while i < len(lines):
        l = lines[i]
        if ANA_MARK.match(l):
            parts = re.split(r"[:：]", l, maxsplit=1)
            if len(parts) == 2 and parts[1].strip():
                return parts[1].strip()
            return " ".join(lines[i + 1:]).strip()
        i += 1
    return ""


def infer_qtype(section: str, options: List[Dict[str, str]], answer: str) -> str:
    if section == "判断题":
        return "judge"
    if options:
        letters = norm_letters(answer)
        if len(letters) >= 2:
            return "multi"
        if len(letters) == 1:
            return "single"
        return "choice"
    if section == "填空题":
        return "fill"
    if section == "简答题":
        return "short"
    if section == "阅读理解":
        return "text"
    return "text"


def parse_block_common(block_text: str, number_in_source: int, section: str, source_name: str) -> List[Dict[str, Any]]:
    """Parse a normal question block (non-reading-subquestion split)."""
    lines = [re.sub(r"\s+", " ", l).strip() for l in block_text.split("\n")]
    lines = [l for l in lines if l]
    if not lines:
        return []

    # remove leading number mark from first line
    first = re.sub(rf"^\s*{number_in_source}[、\.]\s*", "", lines[0]).strip()

    # score like （2）
    score = None
    mscore = re.search(r"（\s*(\d+(?:\.\d+)?)\s*）", first)
    if mscore:
        try:
            score = float(mscore.group(1))
        except Exception:
            score = None
        first = re.sub(r"（\s*\d+(?:\.\d+)?\s*）", "", first).strip()

    stem_parts = [first]
    options: List[Dict[str, str]] = []

    i = 1
    while i < len(lines):
        l = lines[i]
        if ANS_MARK.match(l) or OPT_SINGLE.match(l) or OPT_INLINE.match(l):
            break
        stem_parts.append(l)
        i += 1

    # options
    while i < len(lines):
        l = lines[i]
        if ANS_MARK.match(l):
            break
        m1 = OPT_SINGLE.match(l)
        m2 = OPT_INLINE.match(l)
        if m1:
            key = m1.group(1)
            if i + 1 < len(lines) and not ANS_MARK.match(lines[i + 1]):
                options.append({"key": key, "text": lines[i + 1]})
                i += 2
            else:
                i += 1
            continue
        if m2:
            options.append({"key": m2.group(1), "text": m2.group(2)})
            i += 1
            continue
        i += 1

    answer, next_idx = parse_answer_from_lines(lines, i)
    analysis = parse_analysis_from_lines(lines, next_idx)

    stem = " ".join(stem_parts).strip()
    qtype = infer_qtype(section, options, answer)

    return [
        {
            "source": source_name,
            "number_in_source": number_in_source,
            "section": section,
            "qtype": qtype,
            "stem": stem,
            "options": options,
            "answer": (answer or "").strip(),
            "analysis": (analysis or "").strip(),
            "score": score,
        }
    ]


def split_reading_block(block_text: str, number_in_source: int, source_name: str) -> List[Dict[str, Any]]:
    """Split 阅读理解 into subquestions if (1)(2)... exist."""
    # Remove leading number from the block
    block_wo_num = re.sub(rf"(?m)^\s*{number_in_source}[、\.]\s*", "", block_text, count=1).strip()

    # Find first subquestion marker
    subs = list(SUBQ_START.finditer(block_wo_num))
    if not subs:
        return parse_block_common(block_text, number_in_source, "阅读理解", source_name)

    passage = block_wo_num[: subs[0].start()].strip()
    out: List[Dict[str, Any]] = []

    for idx, m in enumerate(subs):
        sub_no = m.group(1)
        start = m.start()
        end = subs[idx + 1].start() if idx + 1 < len(subs) else len(block_wo_num)
        sub_text = block_wo_num[start:end].strip()

        # Remove leading (k)
        sub_text = re.sub(r"^\s*[（\(]\d+[）\)]\s*", "", sub_text).strip()
        # Parse answer/analysis inside sub_text
        lines = [re.sub(r"\s+", " ", l).strip() for l in sub_text.split("\n")]
        lines = [l for l in lines if l]

        answer, next_idx = parse_answer_from_lines(lines, 0)
        analysis = parse_analysis_from_lines(lines, next_idx)

        # stem is all lines until answer marker
        stem_lines: List[str] = []
        i = 0
        while i < len(lines) and not ANS_MARK.match(lines[i]):
            stem_lines.append(lines[i])
            i += 1

        stem = " ".join(stem_lines).strip()
        full_stem = ""
        if passage:
            full_stem = f"【阅读材料】{passage}\n\n({sub_no}) {stem}".strip()
        else:
            full_stem = f"({sub_no}) {stem}".strip()

        out.append(
            {
                "source": source_name,
                "number_in_source": number_in_source,
                "section": "阅读理解",
                "qtype": "text",
                "stem": full_stem,
                "options": [],
                "answer": (answer or "").strip(),
                "analysis": (analysis or "").strip(),
                "score": None,
            }
        )

    return out


def parse_questions(plain_text: str, source_name: str) -> List[Dict[str, Any]]:
    text = normalize_spaces(plain_text)
    markers = build_section_markers(text)

    ms = list(QSTART.finditer(text))
    blocks: List[Tuple[int, int, int]] = []
    for i, m in enumerate(ms):
        start = m.start()
        end = ms[i + 1].start() if i + 1 < len(ms) else len(text)
        blocks.append((start, end, int(m.group(1))))

    out: List[Dict[str, Any]] = []
    for start, end, number in blocks:
        section = section_at(markers, start)
        block = text[start:end].strip()
        if section == "阅读理解":
            out.extend(split_reading_block(block, number, source_name))
        else:
            out.extend(parse_block_common(block, number, section, source_name))

    return out


# -----------------------------
# Database
# -----------------------------

def init_db(db_path: Path) -> sqlite3.Connection:
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys=ON")

    conn.executescript(
        """
        CREATE TABLE IF NOT EXISTS question(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            qhash TEXT UNIQUE,
            source TEXT,
            number_in_source INTEGER,
            section TEXT,
            qtype TEXT,
            stem TEXT NOT NULL,
            options_json TEXT,
            answer TEXT,
            analysis TEXT,
            score REAL
        );

        CREATE TABLE IF NOT EXISTS attempt(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ts TEXT NOT NULL,
            question_id INTEGER NOT NULL,
            user_answer TEXT,
            is_correct INTEGER,
            mode TEXT,
            FOREIGN KEY(question_id) REFERENCES question(id)
        );

        CREATE TABLE IF NOT EXISTS wrongbook(
            question_id INTEGER PRIMARY KEY,
            added_ts TEXT NOT NULL,
            streak_correct INTEGER NOT NULL DEFAULT 0,
            last_ts TEXT,
            FOREIGN KEY(question_id) REFERENCES question(id)
        );
        """
    )

    # Basic migration for older schema: add qhash/source if missing
    cols = {r[1] for r in conn.execute("PRAGMA table_info(question)").fetchall()}
    if "qhash" not in cols:
        conn.execute("ALTER TABLE question ADD COLUMN qhash TEXT")
    if "source" not in cols:
        conn.execute("ALTER TABLE question ADD COLUMN source TEXT")
    conn.commit()
    return conn


def upsert_questions(conn: sqlite3.Connection, questions: List[Dict[str, Any]]) -> Tuple[int, int]:
    inserted = 0
    skipped = 0
    for q in questions:
        options_json = json.dumps(q.get("options", []), ensure_ascii=False)
        base = json.dumps(
            {
                "stem": q.get("stem", ""),
                "options": q.get("options", []),
                "answer": q.get("answer", ""),
                "section": q.get("section", ""),
            },
            ensure_ascii=False,
            sort_keys=True,
        )
        qhash = sha1_text(base)

        try:
            conn.execute(
                """
                INSERT INTO question(qhash, source, number_in_source, section, qtype, stem, options_json, answer, analysis, score)
                VALUES (?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    qhash,
                    q.get("source"),
                    q.get("number_in_source"),
                    q.get("section"),
                    q.get("qtype"),
                    q.get("stem"),
                    options_json,
                    q.get("answer"),
                    q.get("analysis"),
                    q.get("score"),
                ),
            )
            inserted += 1
        except sqlite3.IntegrityError:
            skipped += 1

    conn.commit()
    return inserted, skipped


def list_question_ids(conn: sqlite3.Connection, mode: str, shuffle: bool) -> List[int]:
    if mode == "wrongbook":
        rows = conn.execute(
            """
            SELECT q.id
            FROM question q
            JOIN wrongbook w ON w.question_id=q.id
            ORDER BY w.last_ts IS NULL DESC, w.last_ts ASC, q.id ASC
            """
        ).fetchall()
        qids = [int(r[0]) for r in rows]
    else:
        rows = conn.execute("SELECT id FROM question ORDER BY id ASC").fetchall()
        qids = [int(r[0]) for r in rows]

    if shuffle and qids:
        # deterministic shuffle per run (no external random state), keep reproducible
        import random

        random.shuffle(qids)
    return qids


def get_question(conn: sqlite3.Connection, qid: int) -> Dict[str, Any]:
    row = conn.execute("SELECT * FROM question WHERE id=?", (qid,)).fetchone()
    if not row:
        raise KeyError(qid)
    d = dict(row)
    d["options"] = json.loads(d["options_json"]) if d.get("options_json") else []
    return d


def record_attempt(conn: sqlite3.Connection, qid: int, user_answer: str, is_correct: Optional[int], mode: str) -> None:
    conn.execute(
        "INSERT INTO attempt(ts, question_id, user_answer, is_correct, mode) VALUES (?,?,?,?,?)",
        (now_ts(), qid, user_answer, is_correct, mode),
    )
    conn.commit()


def fetch_attempts(conn: sqlite3.Connection, qid: int) -> List[sqlite3.Row]:
    return conn.execute(
        "SELECT id, ts, user_answer, is_correct, mode FROM attempt WHERE question_id=? ORDER BY id ASC",
        (qid,),
    ).fetchall()


def fetch_last_attempt(conn: sqlite3.Connection, qid: int, mode: Optional[str] = None) -> Optional[sqlite3.Row]:
    if mode:
        r = conn.execute(
            "SELECT id, ts, user_answer, is_correct, mode FROM attempt WHERE question_id=? AND mode=? ORDER BY id DESC LIMIT 1",
            (qid, mode),
        ).fetchone()
        if r is not None:
            return r
    return conn.execute(
        "SELECT id, ts, user_answer, is_correct, mode FROM attempt WHERE question_id=? ORDER BY id DESC LIMIT 1",
        (qid,),
    ).fetchone()


def compute_unique_accuracy(conn: sqlite3.Connection) -> Tuple[int, int]:
    """Return (correct_unique, judged_unique) by taking the latest judged attempt per question.

    This avoids counting repeated submissions of the same question multiple times in accuracy.
    """
    row = conn.execute(
        """
        SELECT
          SUM(CASE WHEN a.is_correct=1 THEN 1 ELSE 0 END) AS c,
          COUNT(*) AS t
        FROM attempt a
        JOIN (
          SELECT question_id, MAX(id) AS max_id
          FROM attempt
          WHERE is_correct IS NOT NULL
          GROUP BY question_id
        ) m
        ON a.question_id=m.question_id AND a.id=m.max_id
        """
    ).fetchone()
    if not row:
        return 0, 0
    return int(row[0] or 0), int(row[1] or 0)


def compute_total_attempts(conn: sqlite3.Connection) -> int:
    return int(conn.execute("SELECT COUNT(*) FROM attempt").fetchone()[0])


def add_to_wrongbook(conn: sqlite3.Connection, qid: int) -> None:
    ts = now_ts()
    conn.execute(
        """
        INSERT INTO wrongbook(question_id, added_ts, streak_correct, last_ts)
        VALUES (?,?,0,?)
        ON CONFLICT(question_id) DO UPDATE SET last_ts=excluded.last_ts
        """,
        (qid, ts, ts),
    )
    conn.commit()


def remove_from_wrongbook(conn: sqlite3.Connection, qid: int) -> None:
    conn.execute("DELETE FROM wrongbook WHERE question_id=?", (qid,))
    conn.commit()


def update_wrong_streak(conn: sqlite3.Connection, qid: int, is_correct: bool) -> Optional[int]:
    row = conn.execute("SELECT streak_correct FROM wrongbook WHERE question_id=?", (qid,)).fetchone()
    if not row:
        return None
    streak = int(row[0])
    streak = streak + 1 if is_correct else 0
    ts = now_ts()

    if streak >= 3:
        remove_from_wrongbook(conn, qid)
        return 3

    conn.execute("UPDATE wrongbook SET streak_correct=?, last_ts=? WHERE question_id=?", (streak, ts, qid))
    conn.commit()
    return streak


# -----------------------------

# -----------------------------
# GUI (Modern Tkinter/ttkbootstrap)
# -----------------------------

# Optional modern theme. If ttkbootstrap is installed, it will be used automatically.
try:
    import ttkbootstrap as tb  # type: ignore
    HAS_BOOTSTRAP = True
except Exception:
    tb = None  # type: ignore
    HAS_BOOTSTRAP = False

PALETTE = {
    # Light, high-contrast (professional / ins-ish)
    "bg": "#F7F7FC",
    "card": "#FFFFFF",
    "text": "#0F172A",
    "muted": "#475569",
    "border": "#D8DCE6",

    # accents (pastel but vivid)
    "accent": "#FF4D8D",
    "accent2": "#5B5FEF",

    # option states
    "option": "#F3F4F6",
    "option_sel": "#E0E7FF",
    "option_ok": "#BBF7D0",
    "option_bad": "#FECACA",

    # feedback blocks (light but strong)
    "success_bg": "#DCFCE7",
    "success_fg": "#166534",
    "danger_bg": "#FEE2E2",
    "danger_fg": "#991B1B",
    "warn_bg": "#FFEDD5",
    "warn_fg": "#9A3412",

    # badges
    "badge_info_bg": "#E0E7FF",
    "badge_info_fg": "#3730A3",
    "badge_soft_bg": "#ECFDF3",
    "badge_soft_fg": "#14532D",

    "white": "#FFFFFF",
}



def ui_font(size: int, bold: bool = False):
    """Pick a reasonably modern default font across platforms."""
    family = "Segoe UI" if os.name == "nt" else "Arial"
    if bold:
        return (family, size, "bold")
    return (family, size)


BaseWindow = tb.Window if HAS_BOOTSTRAP else tk.Tk  # type: ignore


class QuizApp(BaseWindow):
    def __init__(self, conn: sqlite3.Connection):
        # ttkbootstrap Window needs themename in constructor
        if HAS_BOOTSTRAP:
            super().__init__(themename="flatly")  # type: ignore
        else:
            super().__init__()  # type: ignore

        self.conn = conn

        self.title("本地刷题")
        self.geometry("1024x780")
        self.minsize(900, 650)

        try:
            self.configure(bg=PALETTE["bg"])
        except Exception:
            pass

        # state
        self.mode_var = tk.StringVar(value="normal")  # normal / wrongbook
        self.shuffle_var = tk.BooleanVar(value=False)

        self.qids: List[int] = []
        self.idx = 0
        self.current: Optional[Dict[str, Any]] = None

        # selection state for block-style options
        self.single_selected: str = ""
        self.multi_selected: set[str] = set()
        self.option_buttons: Dict[str, tk.Widget] = {}

        self.feedback_packed = False
        self.stats_window = None

        self._build_ui()

        # auto-import sample if DB is empty
        try:
            total = self.conn.execute("SELECT COUNT(*) FROM question").fetchone()[0]
            sample_doc = APP_DIR / "期末复习题库 (1).doc"
            if int(total) == 0 and sample_doc.exists():
                plain = extract_plain_text(sample_doc)
                qs = parse_questions(plain, sample_doc.name)
                ins, sk = upsert_questions(self.conn, qs)
                messagebox.showinfo(
                    "已自动导入示例题库",
                    f"导入 {ins} 题（重复跳过 {sk} 题）。\n你也可以随时点“导入文档”换成自己的题库。",
                )
        except Exception:
            pass

        self.refresh_question_list(reset_idx=True)
        self.load_current()

        # keyboard shortcuts
        self.bind("<Left>", lambda e: self.prev())
        self.bind("<Right>", lambda e: self.next())
        self.bind("<Return>", lambda e: self.submit())

    # ---------------- UI ----------------
    def _build_ui(self) -> None:
        # Root grid
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(3, weight=1)

        # Header
        header = tk.Frame(self, bg=PALETTE["bg"])
        header.grid(row=0, column=0, sticky="ew", padx=14, pady=(14, 6))
        header.grid_columnconfigure(1, weight=1)

        brand = tk.Frame(header, bg=PALETTE["bg"])
        brand.grid(row=0, column=0, sticky="w")
        tk.Label(
            brand,
            text="Quiz",
            font=ui_font(18, True),
            fg=PALETTE["accent"],
            bg=PALETTE["bg"],
        ).pack(side="left")
        tk.Label(
            brand,
            text=" · 本地刷题",
            font=ui_font(18, True),
            fg=PALETTE["text"],
            bg=PALETTE["bg"],
        ).pack(side="left")

        controls = tk.Frame(header, bg=PALETTE["bg"])
        controls.grid(row=0, column=1, sticky="e")

        ttk.Button(controls, text="导入文档", command=self.on_import).pack(side="left", padx=(0, 10))

        ttk.Label(controls, text="模式：").pack(side="left", padx=(0, 4))
        mode = ttk.Combobox(
            controls,
            width=12,
            state="readonly",
            values=["normal", "wrongbook"],
            textvariable=self.mode_var,
        )
        mode.pack(side="left")
        mode.bind("<<ComboboxSelected>>", lambda e: self.on_mode_change())

        ttk.Checkbutton(
            controls,
            text="随机",
            variable=self.shuffle_var,
            command=lambda: self.refresh_question_list(reset_idx=True),
        ).pack(side="left", padx=10)

        ttk.Label(controls, text="跳到：").pack(side="left")
        self.jump_entry = ttk.Entry(controls, width=8)
        self.jump_entry.pack(side="left")
        ttk.Button(controls, text="跳转", command=self.on_jump).pack(side="left", padx=(6, 0))

        # Accent strip (IG vibe)
        strip = tk.Frame(self, bg=PALETTE["accent"], height=4)
        strip.grid(row=1, column=0, sticky="ew")

        # Status bar (progress + quick stats)
        status = tk.Frame(self, bg=PALETTE["bg"])
        status.grid(row=2, column=0, sticky="ew", padx=14, pady=(10, 6))
        status.grid_columnconfigure(1, weight=1)

        left_status = tk.Frame(status, bg=PALETTE["bg"])
        left_status.grid(row=0, column=0, sticky="w")

        self.progress_text = tk.Label(
            left_status,
            text="",
            font=ui_font(11, True),
            fg=PALETTE["text"],
            bg=PALETTE["bg"],
        )
        self.progress_text.pack(side="left", padx=(0, 10))

        self.progress_var = tk.DoubleVar(value=0.0)
        self.progress = ttk.Progressbar(
            left_status,
            orient="horizontal",
            mode="determinate",
            maximum=100,
            variable=self.progress_var,
            length=360,
        )
        self.progress.pack(side="left")

        right_status = tk.Frame(status, bg=PALETTE["bg"])
        right_status.grid(row=0, column=1, sticky="e")

        def make_badge(text: str, bg: str, fg: str):
            f = tk.Frame(
                right_status,
                bg=bg,
                highlightbackground=PALETTE["border"],
                highlightthickness=1,
            )
            l = tk.Label(f, text=text, font=ui_font(10, True), fg=fg, bg=bg)
            l.pack(padx=10, pady=6)
            f.pack(side="left", padx=6)
            return f, l

        self.badge_accuracy_frame, self.badge_accuracy = make_badge("正确率 --", PALETTE["badge_info_bg"], PALETTE["badge_info_fg"])
        self.badge_wrong_frame, self.badge_wrong = make_badge("错题本 --", PALETTE["danger_bg"], PALETTE["danger_fg"])
        self.badge_dist_frame, self.badge_dist = make_badge("连对分布 --", PALETTE["warn_bg"], PALETTE["warn_fg"])

        # Current wrongbook streak for the current question (only in wrongbook mode)
        self.badge_streak_frame, self.badge_streak = make_badge("", PALETTE["badge_soft_bg"], PALETTE["badge_soft_fg"])
        self.badge_streak_frame.pack_forget()

        self.stats_btn = ttk.Button(right_status, text="统计", command=self.open_stats)
        self.stats_btn.pack(side="left", padx=(10, 0))

        # Main area with scroll
        main = tk.Frame(self, bg=PALETTE["bg"])
        main.grid(row=3, column=0, sticky="nsew", padx=14, pady=10)
        main.grid_columnconfigure(0, weight=1)
        main.grid_rowconfigure(0, weight=1)

        self.canvas = tk.Canvas(main, bg=PALETTE["bg"], highlightthickness=0)
        self.canvas.grid(row=0, column=0, sticky="nsew")
        sb = ttk.Scrollbar(main, orient="vertical", command=self.canvas.yview)
        sb.grid(row=0, column=1, sticky="ns")
        self.canvas.configure(yscrollcommand=sb.set)

        self.inner = tk.Frame(self.canvas, bg=PALETTE["bg"])
        self.canvas_window = self.canvas.create_window((0, 0), window=self.inner, anchor="n")

        def _on_canvas_configure(event):
            # keep a max width to look like a centered card
            w = min(event.width - 10, 940)
            self.canvas.itemconfigure(self.canvas_window, width=w)

        self.canvas.bind("<Configure>", _on_canvas_configure)

        self.inner.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")),
        )

        # Question card
        self.card = tk.Frame(
            self.inner,
            bg=PALETTE["card"],
            highlightbackground=PALETTE["border"],
            highlightthickness=1,
        )
        self.card.pack(fill="x", pady=(0, 14))

        # Card padding wrapper
        pad = tk.Frame(self.card, bg=PALETTE["card"])
        pad.pack(fill="both", expand=True, padx=22, pady=18)

        top_row = tk.Frame(pad, bg=PALETTE["card"])
        top_row.pack(fill="x")
        self.meta_label = tk.Label(
            top_row,
            text="",
            font=ui_font(11, False),
            fg=PALETTE["muted"],
            bg=PALETTE["card"],
        )
        self.meta_label.pack(side="left")

        self.stats_label = tk.Label(
            top_row,
            text="",
            font=ui_font(11, False),
            fg=PALETTE["muted"],
            bg=PALETTE["card"],
        )
        self.stats_label.pack(side="right")

        self.stem_label = tk.Label(
            pad,
            text="",
            font=ui_font(16, True),
            fg=PALETTE["text"],
            bg=PALETTE["card"],
            justify="left",
            wraplength=880,
        )
        self.stem_label.pack(fill="x", pady=(12, 14))

        self.options_frame = tk.Frame(pad, bg=PALETTE["card"])
        self.options_frame.pack(fill="x")

        # Text answer area for subjective questions
        self.text_wrap = tk.Frame(pad, bg=PALETTE["card"])
        self.text_answer = tk.Text(
            self.text_wrap,
            height=6,
            font=ui_font(13, False),
            bd=0,
            highlightthickness=1,
            highlightbackground=PALETTE["border"],
            padx=12,
            pady=10,
            wrap="word",
        )
        self.text_answer.pack(fill="x")
        self.text_wrap.pack(fill="x", pady=(6, 6))

        # Attempt history (only shown after the question has been attempted)
        self.history_box = tk.Frame(pad, bg=PALETTE["card"])

        self.history_row = tk.Frame(self.history_box, bg=PALETTE["card"])
        self.history_row.pack(fill="x")

        self.last_badge = tk.Label(
            self.history_row,
            text="",
            font=ui_font(10, True),
            fg=PALETTE["text"],
            bg=PALETTE["option"],
            padx=10,
            pady=6,
        )
        self.last_badge.pack(side="left")

        self.history_seq = tk.Label(
            self.history_row,
            text="",
            font=ui_font(11, False),
            fg=PALETTE["muted"],
            bg=PALETTE["card"],
        )
        self.history_seq.pack(side="left", padx=(10, 0))

        self.history_detail = tk.Label(
            self.history_box,
            text="",
            font=ui_font(11, False),
            fg=PALETTE["muted"],
            bg=PALETTE["card"],
            justify="left",
            wraplength=880,
        )
        self.history_detail.pack(anchor="w", pady=(6, 0))

        self.history_box.pack_forget()

        # Feedback block (colored)
        self.feedback = tk.Frame(
            pad,
            bg=PALETTE["option"],
            highlightbackground=PALETTE["border"],
            highlightthickness=1,
        )
        self.feedback_title = tk.Label(
            self.feedback,
            text="",
            font=ui_font(12, True),
            fg=PALETTE["text"],
            bg=PALETTE["option"],
        )
        self.feedback_title.pack(anchor="w", padx=14, pady=(10, 2))
        self.feedback_label = tk.Label(
            self.feedback,
            text="",
            font=ui_font(12, False),
            fg=PALETTE["text"],
            bg=PALETTE["option"],
            justify="left",
            wraplength=880,
        )
        self.feedback_label.pack(fill="x", padx=14, pady=(0, 12))

        # Footer actions
        footer = tk.Frame(self, bg=PALETTE["bg"])
        footer.grid(row=4, column=0, sticky="ew", padx=14, pady=(0, 14))
        footer.grid_columnconfigure(1, weight=1)

        left = tk.Frame(footer, bg=PALETTE["bg"])
        left.grid(row=0, column=0, sticky="w")
        ttk.Button(left, text="上一题", command=self.prev).pack(side="left")
        ttk.Button(left, text="提交", command=self.submit).pack(side="left", padx=8)
        ttk.Button(left, text="下一题", command=self.next).pack(side="left")

        right = tk.Frame(footer, bg=PALETTE["bg"])
        right.grid(row=0, column=2, sticky="e")
        ttk.Button(right, text="手动标记正确", command=lambda: self.mark_subjective(True)).pack(side="left", padx=(0, 8))
        ttk.Button(right, text="手动标记错误", command=lambda: self.mark_subjective(False)).pack(side="left")

    # ---------------- state helpers ----------------
    def update_stats(self) -> None:
        # overall stats
        total_all = int(self.conn.execute("SELECT COUNT(*) FROM question").fetchone()[0])
        wrong = int(self.conn.execute("SELECT COUNT(*) FROM wrongbook").fetchone()[0])

        correct, judged = compute_unique_accuracy(self.conn)
        total_attempts = compute_total_attempts(self.conn)
        acc = (correct / judged * 100.0) if judged else None

        # streak distribution in wrongbook
        dist = {0: 0, 1: 0, 2: 0}
        for r in self.conn.execute("SELECT streak_correct, COUNT(*) FROM wrongbook GROUP BY streak_correct").fetchall():
            k = int(r[0])
            if k in dist:
                dist[k] = int(r[1])

        # card top-right (compact)
        acc_text = f"正确率 {acc:.0f}%" if acc is not None else "正确率 --"
        self.stats_label.config(text=f"{acc_text} · 错题本 {wrong} · 已做 {judged} · 作答 {total_attempts}")

        # status bar badges
        self.badge_accuracy.config(text=acc_text)
        self.badge_wrong.config(text=f"错题本 {wrong}")
        self.badge_dist.config(text=f"连对分布 0:{dist[0]} 1:{dist[1]} 2:{dist[2]}")

        # progress
        total_mode = len(self.qids) if self.qids else 0
        cur = (self.idx + 1) if self.qids else 0
        if total_mode:
            pct = cur / total_mode * 100.0
            self.progress_var.set(pct)
            self.progress_text.config(text=f"进度 {cur}/{total_mode}")
        else:
            self.progress_var.set(0.0)
            self.progress_text.config(text="进度 0/0")

    def refresh_question_list(self, reset_idx: bool) -> None:
        self.qids = list_question_ids(self.conn, self.mode_var.get(), self.shuffle_var.get())
        if reset_idx:
            self.idx = 0
        self.update_stats()

    def on_mode_change(self) -> None:
        self.refresh_question_list(reset_idx=True)
        self.load_current()

    def on_jump(self) -> None:
        if not self.qids:
            return
        try:
            k = int(self.jump_entry.get().strip())
        except Exception:
            return
        k = max(1, min(len(self.qids), k))
        self.idx = k - 1
        self.load_current()

    def on_import(self) -> None:
        path = filedialog.askopenfilename(
            title="选择题库文件",
            filetypes=[
                ("Documents", "*.docx *.doc *.html *.htm *.txt *.md"),
                ("All", "*.*"),
            ],
        )
        if not path:
            return
        f = Path(path)
        try:
            plain = extract_plain_text(f)
            qs = parse_questions(plain, source_name=f.name)
            ins, skip = upsert_questions(self.conn, qs)
        except Exception as e:
            messagebox.showerror("导入失败", str(e))
            return

        messagebox.showinfo("导入完成", f"识别到 {len(qs)} 道题\n新增 {ins}，去重跳过 {skip}")
        self.refresh_question_list(reset_idx=True)
        self.load_current()

    # ---------------- rendering ----------------
    def clear_options(self) -> None:
        for w in self.options_frame.winfo_children():
            w.destroy()
        self.single_selected = ""
        self.multi_selected.clear()
        self.option_buttons.clear()

    def _format_attempt_history(self, attempts: List[sqlite3.Row]) -> Tuple[str, str]:
        """Return (seq_text, detail_text) for a question."""
        judged = [r for r in attempts if r[3] is not None]  # is_correct
        # sequence like: 正确-错误-正确
        seq_parts: List[str] = []
        for r in judged[-20:]:
            seq_parts.append("正确" if int(r[3]) == 1 else "错误")
        seq_text = "历史：" + ("-".join(seq_parts) if seq_parts else "—")

        # last attempt (prefer current mode, fallback to any)
        last = None
        cur_mode = self.mode_var.get()
        for r in reversed(attempts):
            if r[4] == cur_mode:
                last = r
                break
        if last is None and attempts:
            last = attempts[-1]
        if not last:
            return seq_text, ""

        ua = (last[2] or "").strip().replace("\n", " ")
        if len(ua) > 90:
            ua = ua[:90] + "…"

        res = last[3]
        if res is None:
            res_text = "未判定"
        else:
            res_text = "正确" if int(res) == 1 else "错误"
        detail = f"上次：{res_text} · 你的答案：{ua if ua else '(空)'}"
        return seq_text, detail

    def _set_last_badge(self, is_correct: Optional[int]) -> None:
        if is_correct is None:
            bg, fg, txt = PALETTE["warn_bg"], PALETTE["warn_fg"], "上次未判定"
        elif int(is_correct) == 1:
            bg, fg, txt = PALETTE["success_bg"], PALETTE["success_fg"], "上次正确"
        else:
            bg, fg, txt = PALETTE["danger_bg"], PALETTE["danger_fg"], "上次错误"
        self.last_badge.configure(text=txt, bg=bg, fg=fg)

    def _restore_previous_answer(self, q: Dict[str, Any]) -> None:
        """Restore last answer (prefer current mode) so going back keeps selection/history."""
        qid = int(q["id"])
        last = fetch_last_attempt(self.conn, qid, mode=self.mode_var.get())
        if last is None:
            return

        user_answer = (last[2] or "").strip()
        if q.get("options"):
            if (q.get("qtype") or "").lower() == "multi":
                self.multi_selected = set(norm_letters(user_answer))
            else:
                self.single_selected = norm_letters(user_answer)[:1] if user_answer else ""
            self._refresh_selection_styles()
        else:
            # subjective: restore previous input
            self.text_answer.delete("1.0", "end")
            if user_answer:
                self.text_answer.insert("1.0", user_answer)

    def _render_history(self, q: Dict[str, Any]) -> None:
        qid = int(q["id"])
        attempts = fetch_attempts(self.conn, qid)
        if not attempts:
            self.history_box.pack_forget()
            return

        # badge: last attempt in current mode preferred
        last = fetch_last_attempt(self.conn, qid, mode=self.mode_var.get())
        if last is None:
            last = attempts[-1]
        self._set_last_badge(last[3])

        seq_text, detail = self._format_attempt_history(attempts)
        self.history_seq.configure(text=seq_text)
        self.history_detail.configure(text=detail)

        if not self.history_box.winfo_ismapped():
            self.history_box.pack(fill="x", pady=(10, 0))

    def _set_feedback_style(self, kind: str, title: str, body: str) -> None:
        if kind == "success":
            bg = PALETTE["success_bg"]
            fg = PALETTE["success_fg"]
        elif kind == "danger":
            bg = PALETTE["danger_bg"]
            fg = PALETTE["danger_fg"]
        elif kind == "warn":
            bg = PALETTE["warn_bg"]
            fg = PALETTE["warn_fg"]
        else:
            bg = PALETTE["card"]
            fg = PALETTE["text"]

        self.feedback.configure(
            bg=bg,
            highlightbackground=PALETTE["border"],
            highlightthickness=1,
        )
        self.feedback_title.configure(text=title, bg=bg, fg=fg)
        self.feedback_label.configure(text=body, bg=bg, fg=fg)

    def _hide_feedback(self) -> None:
        if getattr(self, "feedback_packed", False):
            self.feedback.pack_forget()
            self.feedback_packed = False
        self._set_feedback_style("neutral", "", "")

    def _show_feedback(self, kind: str, title: str, body: str) -> None:
        if not getattr(self, "feedback_packed", False):
            self.feedback.pack(fill="x", pady=(14, 0))
            self.feedback_packed = True
        self._set_feedback_style(kind, title, body)

    def _style_option_button(self, key: str, style: str) -> None:
        btn = self.option_buttons.get(key)
        if not btn:
            return

        # using plain tk.Button for full color control
        if style == "neutral":
            btn.configure(bg=PALETTE["option"], fg=PALETTE["text"], activebackground=PALETTE["option_sel"], activeforeground=PALETTE["text"])
        elif style == "selected":
            btn.configure(bg=PALETTE["option_sel"], fg=PALETTE["text"], activebackground=PALETTE["option_sel"], activeforeground=PALETTE["text"])
        elif style == "success":
            btn.configure(bg=PALETTE["option_ok"], fg=PALETTE["text"], activebackground=PALETTE["option_ok"], activeforeground=PALETTE["text"])
        elif style == "danger":
            btn.configure(bg=PALETTE["option_bad"], fg=PALETTE["text"], activebackground=PALETTE["option_bad"], activeforeground=PALETTE["text"])

    def _refresh_selection_styles(self) -> None:
        if not self.current:
            return
        opts = self.current.get("options") or []
        if not opts:
            return

        qtype = self.current.get("qtype")
        for opt in opts:
            k = opt.get("key", "")
            if qtype == "multi":
                self._style_option_button(k, "selected" if k in self.multi_selected else "neutral")
            else:
                self._style_option_button(k, "selected" if k == self.single_selected else "neutral")

    def _make_option_button(self, key: str, text: str, qtype: str) -> None:
        label = f"{key}. {text}".strip()

        def on_click():
            if qtype == "multi":
                if key in self.multi_selected:
                    self.multi_selected.remove(key)
                else:
                    self.multi_selected.add(key)
            else:
                self.single_selected = key
            self._refresh_selection_styles()

        btn = tk.Button(
            self.options_frame,
            text=label,
            command=on_click,
            font=ui_font(15, False),
            bg=PALETTE["option"],
            fg=PALETTE["text"],
            activebackground=PALETTE["option_sel"],
            relief="flat",
            bd=0,
            padx=14,
            pady=12,
            anchor="w",
            highlightthickness=1,
            highlightbackground=PALETTE["border"],
            justify="left",
            wraplength=880,
            cursor="hand2",
        )
        btn.pack(fill="x", pady=6)
        self.option_buttons[key] = btn

    def load_current(self) -> None:
        self.clear_options()
        self.text_answer.delete("1.0", "end")
        self._hide_feedback()
        self.history_box.pack_forget()

        if not self.qids:
            self.meta_label.config(text="还没有题目。点击“导入文档”开始。")
            self.stem_label.config(text="")
            self.text_wrap.pack_forget()
            return

        qid = self.qids[self.idx]
        q = get_question(self.conn, qid)
        self.current = q

        meta = f"[{self.mode_var.get()}]  第 {self.idx + 1}/{len(self.qids)}  ·  {q.get('section','')}  ·  来源：{q.get('source','')}"
        self.meta_label.config(text=meta)
        self.stem_label.config(text=q.get("stem", ""))

        opts = q.get("options") or []
        qtype = q.get("qtype") or "text"

        if opts:
            self.text_wrap.pack_forget()
            for opt in opts:
                self._make_option_button(opt.get("key", ""), opt.get("text", ""), qtype)
            # restore last selection (if any)
            self._restore_previous_answer(q)
            self._refresh_selection_styles()
        else:
            self.text_wrap.pack(fill="x", pady=(6, 6))
            self._restore_previous_answer(q)

        self._render_history(q)

        # update streak badge
        if self.mode_var.get() == "wrongbook":
            row = self.conn.execute("SELECT streak_correct FROM wrongbook WHERE question_id=?", (qid,)).fetchone()
            if row is not None:
                self.badge_streak.config(text=f"本题连对 {int(row[0])}/3")
                self.badge_streak_frame.pack(side="left", padx=6, before=self.stats_btn)
            else:
                self.badge_streak_frame.pack_forget()
        else:
            self.badge_streak_frame.pack_forget()

        self.canvas.yview_moveto(0.0)
        self.update_stats()

    # ---------------- navigation ----------------
    def prev(self) -> None:
        if not self.qids:
            return
        self.idx = max(0, self.idx - 1)
        self.load_current()

    def next(self) -> None:
        if not self.qids:
            return
        self.idx = min(len(self.qids) - 1, self.idx + 1)
        self.load_current()

    # ---------------- answer logic ----------------
    def _get_user_answer(self, q: Dict[str, Any]) -> Tuple[str, Optional[int]]:
        qtype = q.get("qtype")
        answer_std = (q.get("answer") or "").strip()

        if q.get("options"):
            if qtype == "multi":
                user = "".join(sorted(self.multi_selected))
                a = norm_letters(answer_std)
                u = norm_letters(user)
                if a:
                    return user, (1 if u == a else 0)
                return user, None
            else:
                user = (self.single_selected or "").strip()
                a = norm_letters(answer_std)
                u = norm_letters(user)
                if a:
                    return user, (1 if u == a else 0)
                return user, None

        user = self.text_answer.get("1.0", "end").strip()
        if not answer_std:
            return user, None

        if qtype == "judge":
            b1 = judge_to_bool(user)
            b2 = judge_to_bool(answer_std)
            if b1 is not None and b2 is not None:
                return user, (1 if b1 == b2 else 0)
            return user, None

        if qtype == "fill":
            u = re.sub(r"\s+", "", user)
            a = re.sub(r"\s+", "", answer_std)
            return user, (1 if u == a else 0)

        u = re.sub(r"\s+", " ", user).strip().casefold()
        a = re.sub(r"\s+", " ", answer_std).strip().casefold()
        return user, (1 if u == a else 0)

    def _apply_result_styles(self, q: Dict[str, Any], user_answer: str, is_correct: Optional[int]) -> None:
        # Color option blocks if it is an option question and standard answer exists
        opts = q.get("options") or []
        if not opts:
            return

        correct = norm_letters(q.get("answer") or "")
        user = norm_letters(user_answer)
        if not correct:
            return

        for opt in opts:
            k = opt.get("key", "")
            if k in correct:
                self._style_option_button(k, "success")
            elif k in user:
                self._style_option_button(k, "danger")
            else:
                self._style_option_button(k, "neutral")

    def submit(self) -> None:
        if not self.current:
            return
        q = self.current
        qid = q["id"]
        mode = self.mode_var.get()

        user_answer, is_correct = self._get_user_answer(q)
        record_attempt(self.conn, qid, user_answer, is_correct, mode)

        # wrongbook rules
        if mode == "normal":
            if is_correct == 0:
                add_to_wrongbook(self.conn, qid)
        else:
            if is_correct is not None:
                streak = update_wrong_streak(self.conn, qid, bool(is_correct))
                if streak == 3:
                    messagebox.showinfo("移除错题", "该题已在错题练习中连续答对3次，已移除错题本")
                    self.refresh_question_list(reset_idx=False)
                    if self.idx >= len(self.qids):
                        self.idx = max(0, len(self.qids) - 1)

        # feedback text
        body = f"你的答案：{user_answer if user_answer else '(空)'}\n"
        body += f"标准答案：{q.get('answer','')}\n"
        if q.get("analysis"):
            body += f"解析：{q.get('analysis')}\n"

        if is_correct is None:
            title = "🤔 无法自动判定"
            body += "（主观题/无标准答案）可用下方按钮手动标记。"
            self._show_feedback("warn", title, body)
        elif is_correct == 1:
            title = "✅ 正确"
            self._show_feedback("success", title, body)
        else:
            title = "❌ 错误"
            self._show_feedback("danger", title, body)

        self._apply_result_styles(q, user_answer, is_correct)
        self._render_history(q)
        self.update_stats()

    def mark_subjective(self, correct: bool) -> None:
        if not self.current:
            return
        q = self.current
        qid = q["id"]
        mode = self.mode_var.get()
        user = self.text_answer.get("1.0", "end").strip()

        record_attempt(self.conn, qid, user, 1 if correct else 0, mode)

        if mode == "normal":
            if not correct:
                add_to_wrongbook(self.conn, qid)
        else:
            streak = update_wrong_streak(self.conn, qid, correct)
            if streak == 3:
                messagebox.showinfo("移除错题", "该题已在错题练习中连续答对3次，已移除错题本")
                self.refresh_question_list(reset_idx=False)
                if self.idx >= len(self.qids):
                    self.idx = max(0, len(self.qids) - 1)

        if correct:
            self._show_feedback("success", "✅ 已标记正确", "已记录本次作答。")
        else:
            self._show_feedback("danger", "❌ 已标记错误", "已记录本次作答，并加入错题本。")
        self._render_history(q)
        self.update_stats()

    def open_stats(self) -> None:
        # single instance
        if self.stats_window is not None and tk.Toplevel.winfo_exists(self.stats_window):
            self.stats_window.lift()
            return

        win = tk.Toplevel(self)
        self.stats_window = win
        win.title("统计")
        win.geometry("720x520")
        try:
            win.configure(bg=PALETTE["bg"])
        except Exception:
            pass

        # compute stats
        total_all = int(self.conn.execute("SELECT COUNT(*) FROM question").fetchone()[0])
        wrong = int(self.conn.execute("SELECT COUNT(*) FROM wrongbook").fetchone()[0])
        correct, judged = compute_unique_accuracy(self.conn)
        total_attempts = compute_total_attempts(self.conn)
        acc = (correct / judged * 100.0) if judged else None

        dist = {0: 0, 1: 0, 2: 0}
        for r in self.conn.execute("SELECT streak_correct, COUNT(*) FROM wrongbook GROUP BY streak_correct").fetchall():
            k = int(r[0])
            if k in dist:
                dist[k] = int(r[1])

        # wrongbook by section (top 6)
        sec_rows = self.conn.execute(
            """
            SELECT q.section, COUNT(*) AS n
            FROM question q
            JOIN wrongbook w ON w.question_id=q.id
            GROUP BY q.section
            ORDER BY n DESC
            LIMIT 6
            """
        ).fetchall()

        # card
        card = tk.Frame(win, bg=PALETTE["card"], highlightbackground=PALETTE["border"], highlightthickness=1)
        card.pack(fill="both", expand=True, padx=18, pady=18)

        pad = tk.Frame(card, bg=PALETTE["card"])
        pad.pack(fill="both", expand=True, padx=18, pady=16)

        tk.Label(pad, text="学习统计", font=ui_font(16, True), fg=PALETTE["text"], bg=PALETTE["card"]).pack(anchor="w")
        tk.Label(
            pad,
            text=f"总题数：{total_all}   已判定作答：{judged}   错题本：{wrong}",
            font=ui_font(12, False),
            fg=PALETTE["muted"],
            bg=PALETTE["card"],
        ).pack(anchor="w", pady=(6, 10))

        acc_text = f"{acc:.1f}%" if acc is not None else "--"
        kpi = tk.Frame(pad, bg=PALETTE["card"])
        kpi.pack(fill="x", pady=(0, 10))

        def kpi_block(title, value, bg, fg):
            f = tk.Frame(kpi, bg=bg, highlightbackground=PALETTE["border"], highlightthickness=1)
            tk.Label(f, text=title, font=ui_font(10, True), fg=fg, bg=bg).pack(anchor="w", padx=12, pady=(10, 0))
            tk.Label(f, text=value, font=ui_font(18, True), fg=fg, bg=bg).pack(anchor="w", padx=12, pady=(2, 10))
            f.pack(side="left", padx=6, fill="x", expand=True)

        kpi_block("正确率", acc_text, PALETTE["badge_info_bg"], PALETTE["badge_info_fg"])
        kpi_block("错题本", str(wrong), PALETTE["danger_bg"], PALETTE["danger_fg"])
        kpi_block("已做题", str(judged), PALETTE["warn_bg"], PALETTE["warn_fg"])
        kpi_block("作答次数", str(total_attempts), PALETTE["option_sel"], PALETTE["text"])

        # distribution bar
        tk.Label(pad, text="错题连对分布（0/1/2）", font=ui_font(12, True), fg=PALETTE["text"], bg=PALETTE["card"]).pack(anchor="w", pady=(6, 6))
        bar = tk.Canvas(pad, height=18, bg=PALETTE["card"], highlightthickness=1, highlightbackground=PALETTE["border"])
        bar.pack(fill="x")
        total_dist = dist[0] + dist[1] + dist[2]
        w = 640
        bar.configure(width=w)
        x = 0
        def seg(n, color):
            nonlocal x
            if total_dist == 0:
                return
            segw = int(w * (n / total_dist))
            bar.create_rectangle(x, 0, x + segw, 18, fill=color, outline=color)
            x += segw
        # light red / amber / light green
        seg(dist[0], PALETTE["danger_bg"])
        seg(dist[1], PALETTE["warn_bg"])
        seg(dist[2], PALETTE["success_bg"])
        tk.Label(pad, text=f"0连对：{dist[0]}    1连对：{dist[1]}    2连对：{dist[2]}", font=ui_font(11, False), fg=PALETTE["muted"], bg=PALETTE["card"]).pack(anchor="w", pady=(6, 10))

        # wrongbook sections
        tk.Label(pad, text="错题按章节（Top）", font=ui_font(12, True), fg=PALETTE["text"], bg=PALETTE["card"]).pack(anchor="w", pady=(6, 6))
        if sec_rows:
            for r in sec_rows:
                tk.Label(pad, text=f"• {r[0]}：{int(r[1])}", font=ui_font(11, False), fg=PALETTE["text"], bg=PALETTE["card"]).pack(anchor="w")
        else:
            tk.Label(pad, text="（当前错题本为空）", font=ui_font(11, False), fg=PALETTE["muted"], bg=PALETTE["card"]).pack(anchor="w")

        ttk.Button(pad, text="关闭", command=win.destroy).pack(anchor="e", pady=(12, 0))


def main() -> None:
    conn = init_db(DB_PATH)
    app = QuizApp(conn)
    app.mainloop()


if __name__ == "__main__":
    os.chdir(str(APP_DIR))
    main()
